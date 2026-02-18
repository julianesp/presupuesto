"""Genera el manual de operacion del Sistema Presupuestal en formato Word (.docx)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ============ ESTILOS ============
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0, 51, 102)

# Margenes
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_table(headers, rows, col_widths=None):
    """Agrega tabla con formato."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Encabezados
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Datos
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def p(text, bold=False, italic=False, size=11, color=None, align=None):
    """Agrega parrafo con formato."""
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align == 'center':
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return par


def nota(text):
    """Agrega nota destacada."""
    par = doc.add_paragraph()
    run = par.add_run("NOTA: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(153, 51, 0)
    run2 = par.add_run(text)
    run2.font.size = Pt(10)
    run2.italic = True


def importante(text):
    """Agrega texto importante."""
    par = doc.add_paragraph()
    run = par.add_run("IMPORTANTE: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(192, 0, 0)
    run2 = par.add_run(text)
    run2.font.size = Pt(10)
    run2.bold = True


def paso(numero, texto):
    """Agrega un paso numerado."""
    par = doc.add_paragraph()
    run = par.add_run(f"PASO {numero}: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 102, 51)
    run2 = par.add_run(texto)
    run2.font.size = Pt(11)


def bullet(text):
    par = doc.add_paragraph(text, style='List Bullet')
    for run in par.runs:
        run.font.size = Pt(10)
    return par


def numbered(text):
    par = doc.add_paragraph(text, style='List Number')
    for run in par.runs:
        run.font.size = Pt(10)
    return par


def ruta(text):
    """Muestra ruta de menu."""
    par = doc.add_paragraph()
    run = par.add_run("Ruta: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = par.add_run(text)
    run2.font.size = Pt(10)
    run2.italic = True
    run2.font.color.rgb = RGBColor(0, 102, 153)


def ejemplo(titulo, lineas):
    """Ejemplo practico en recuadro."""
    par = doc.add_paragraph()
    run = par.add_run(f"Ejemplo: {titulo}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(102, 0, 102)
    for linea in lineas:
        par2 = doc.add_paragraph()
        par2.paragraph_format.left_indent = Cm(1)
        run2 = par2.add_run(linea)
        run2.font.size = Pt(10)
        run2.font.name = 'Consolas'


# =====================================================
# PORTADA
# =====================================================
doc.add_paragraph()
doc.add_paragraph()
p("MANUAL DE OPERACION", bold=True, size=26, color=(0, 51, 102), align='center')
p("SISTEMA PRESUPUESTAL IE 2026", bold=True, size=20, color=(0, 51, 102), align='center')
doc.add_paragraph()
p("Institucion Educativa Almirante Padilla", bold=True, size=14, align='center')
p("Putumayo - Colombia", size=12, align='center')
doc.add_paragraph()
p("Codigo DANE: 186755000015", size=12, align='center')
doc.add_paragraph()
doc.add_paragraph()
p("Version 1.0 - Febrero 2026", italic=True, size=11, align='center')

doc.add_page_break()

# =====================================================
# TABLA DE CONTENIDO
# =====================================================
doc.add_heading('TABLA DE CONTENIDO', level=1)
contenido = [
    "1. Inicio del Sistema y Dashboard",
    "2. Configuracion Inicial",
    "3. Plan Presupuestal - Rubros",
    "4. Gestion de Terceros",
    "5. Registro de Movimiento - Cadena Presupuestal",
    "    5.1 CDP (Certificado de Disponibilidad Presupuestal)",
    "    5.2 RP (Registro Presupuestal / Compromiso)",
    "    5.3 Obligacion",
    "    5.4 Pago",
    "    5.5 Recaudo (Ingresos)",
    "6. Modificaciones Presupuestales",
    "7. Consultas y Listados",
    "8. Generacion de Informes",
    "9. Informe SIFSE (Trimestral)",
    "10. Mapeo SIFSE",
    "11. Plan Anualizado de Caja (PAC)",
    "12. Procesos Presupuestales",
    "13. Copias de Seguridad",
    "14. Guia Rapida de Pruebas",
]
for c in contenido:
    par = doc.add_paragraph(c)
    par.paragraph_format.space_after = Pt(2)
    for run in par.runs:
        run.font.size = Pt(11)

doc.add_page_break()

# =====================================================
# 1. INICIO DEL SISTEMA
# =====================================================
doc.add_heading('1. Inicio del Sistema y Dashboard', level=1)

doc.add_heading('Como iniciar', level=2)
p("Ejecute el archivo main.py o el acceso directo del sistema. La ventana principal se abrira maximizada.")

doc.add_heading('Pantalla principal (Dashboard)', level=2)
p("Al iniciar vera la pantalla principal con la siguiente informacion:")

p("Encabezado:", bold=True)
bullet("Nombre de la IE, NIT, Vigencia, Mes activo, Rector, Tesorero")

p("Botones de acceso rapido organizados en 4 secciones:", bold=True)
add_table(
    ["Seccion", "Botones disponibles"],
    [
        ["Plan Presupuestal", "Rubros Gastos, Rubros Ingresos, Resumen, Terceros, PAC"],
        ["Registro de Movimiento", "1.CDP, 2.RP, 3.OBLIGACION, 4.PAGO, 5.RECAUDO"],
        ["Modificaciones", "6.ADICION, 7.REDUCCION, 8.CRED/CONTRACRED"],
        ["Informes y Procesos", "INFORMES, TARJETAS, LISTADOS, CONSOLIDAR, CIERRE MES"],
    ],
    col_widths=[5, 12]
)

p("Panel de resumen financiero:", bold=True)
bullet("Izquierda: Gastos (Apropiacion, CDPs, Compromisos, Obligaciones, Pagos, Saldos)")
bullet("Derecha: Ingresos (Presupuesto, Recaudado, Saldo por Recaudar)")
bullet("Indicador de equilibrio presupuestal (verde = equilibrado, rojo = desbalance)")

doc.add_heading('Barra de menu superior', level=2)
add_table(
    ["Menu", "Opciones principales"],
    [
        ["Plan Presupuestal", "Rubros Gastos, Rubros Ingresos, Importar Excel/CSV, Resumen, PAC"],
        ["Gestion de Terceros", "Registrar Tercero, Ver Terceros"],
        ["Registro de Movimiento", "CDP, RP, Obligacion, Pago, Recaudo, Adicion, Reduccion, Credito/Contracredito"],
        ["Consultas", "Ver CDPs/RPs/Obligaciones/Pagos/Recaudos, Tarjeta, Modificaciones"],
        ["Generacion de Informes", "Informes y Ejecuciones, Informe SIFSE (Trimestral)"],
        ["Procesos Presupuestales", "Consolidar Mes, Cierre de Mes"],
        ["Configuracion", "Config General, Mapeo SIFSE, Backup, Restaurar, Salir"],
    ],
    col_widths=[5, 12]
)

doc.add_page_break()

# =====================================================
# 2. CONFIGURACION INICIAL
# =====================================================
doc.add_heading('2. Configuracion Inicial', level=1)
ruta("Menu Configuracion > Configuracion General")
p("Antes de empezar a usar el sistema, configure estos datos:")

add_table(
    ["Campo", "Que poner", "Ejemplo"],
    [
        ["Vigencia", "Ano fiscal", "2026"],
        ["Institucion Educativa", "Nombre completo de la IE", "IE Almirante Padilla"],
        ["NIT Institucion", "NIT con DV", "800.123.456-7"],
        ["Codigo DANE", "Codigo del MEN (12 digitos)", "186755000015"],
        ["Rector/Director", "Nombre del rector", "Juan Perez Garcia"],
        ["Tesorero/Pagador", "Nombre de la tesorera", "Martha Riobamba"],
        ["Mes Actual", "Mes en curso (1-12)", "2"],
    ],
    col_widths=[4, 5.5, 5.5]
)

p('Haga clic en "Guardar" para aplicar los cambios.')
importante("El Codigo DANE es necesario para el informe SIFSE. Viene precargado con 186755000015.")

doc.add_page_break()

# =====================================================
# 3. PLAN PRESUPUESTAL
# =====================================================
doc.add_heading('3. Plan Presupuestal - Rubros', level=1)

doc.add_heading('3.1 Importar el catalogo desde Excel', level=2)
ruta("Menu Plan Presupuestal > Importar Catalogo Excel")
numbered('Haga clic en "Importar Catalogo Excel"')
numbered("Seleccione el archivo .xlsx o .xlsm con el presupuesto aprobado")
numbered("El sistema importara rubros de gastos e ingresos automaticamente")
numbered("Mostrara: cantidad importada, totales, y si hay equilibrio presupuestal")

doc.add_heading('3.2 Importar desde CSV (archivo plano)', level=2)
ruta("Menu Plan Presupuestal > Importar Archivo Plano")
p("Formato requerido: Separador punto y coma (;), codificacion UTF-8. La primera fila (encabezados) se ignora.")

add_table(
    ["Tipo", "Columnas del archivo"],
    [
        ["Rubros de Gastos", "codigo;cuenta;apropiacion_inicial"],
        ["Rubros de Ingresos", "codigo;cuenta;presupuesto_inicial"],
        ["Terceros", "nit;dv;nombre;direccion;telefono;email;tipo;banco;tipo_cuenta;no_cuenta"],
        ["Conceptos", "codigo_rubro;concepto"],
    ],
    col_widths=[4.5, 12]
)
nota("Puede descargar plantillas CSV desde el mismo submenu para ver el formato exacto.")

doc.add_heading('3.3 Gestion de Rubros de Gastos', level=2)
ruta('Menu Plan Presupuestal > Gestion Rubros de Gastos (o boton "Rubros de Gastos" en dashboard)')
bullet("Barra de busqueda: filtra por codigo o nombre en tiempo real")
bullet("Tabla con: Codigo, Cuenta, Hoja (Si/No), Aprop. Inicial, Aprop. Definitiva")
bullet("Crear Rubro: Ingrese codigo, nombre y apropiacion")
bullet("Editar Rubro: Seleccione y modifique nombre o valores")
bullet("Eliminar Rubro: Solo si no tiene CDPs ni rubros hijos")
importante("Solo los rubros HOJA pueden tener movimientos (CDPs, RPs, etc.). Los rubros padre se calculan automaticamente sumando sus hijos.")

doc.add_heading('3.4 Gestion de Rubros de Ingresos', level=2)
p("Funciona igual que gastos. Los campos son Presupuesto Inicial y Presupuesto Definitivo.")

doc.add_heading('3.5 Resumen Presupuestal', level=2)
ruta("Menu Plan Presupuestal > Resumen Presupuestal")
bullet("Panel izquierdo: Arbol de rubros (seleccione Gastos o Ingresos)")
bullet("Panel derecho: Detalle del rubro seleccionado con todos los valores")
bullet("Los valores son enlaces clickeables: haga clic para ver los documentos asociados")

doc.add_page_break()

# =====================================================
# 4. TERCEROS
# =====================================================
doc.add_heading('4. Gestion de Terceros', level=1)

doc.add_heading('4.1 Registrar un tercero', level=2)
ruta("Menu Gestion de Terceros > Registrar Tercero")

add_table(
    ["Campo", "Obligatorio", "Descripcion"],
    [
        ["NIT/CC", "Si", "Cedula o NIT del proveedor"],
        ["DV", "No", "Digito de verificacion"],
        ["Nombre/Razon Social", "Si", "Nombre completo o razon social"],
        ["Direccion", "No", "Direccion del tercero"],
        ["Telefono", "No", "Numero de contacto"],
        ["Email", "No", "Correo electronico"],
        ["Tipo", "Si", "Natural o Juridica (lista desplegable)"],
        ["Banco", "No", "Entidad bancaria para pagos"],
        ["Tipo Cuenta", "No", "Ahorros o Corriente"],
        ["No. Cuenta", "No", "Numero de cuenta bancaria"],
    ],
    col_widths=[4, 2.5, 9]
)

doc.add_heading('4.2 Ver Terceros', level=2)
ruta("Menu Gestion de Terceros > Ver Terceros")
p("Muestra la lista completa con NIT, Nombre, Tipo, Telefono y Email.")
nota("Tambien puede crear terceros rapidos al momento de registrar un RP. El sistema le preguntara si desea registrarlo.")

doc.add_page_break()

# =====================================================
# 5. CADENA PRESUPUESTAL
# =====================================================
doc.add_heading('5. Registro de Movimiento - Cadena Presupuestal', level=1)

p("La cadena presupuestal sigue un orden obligatorio:", bold=True)
p("CDP  -->  RP  -->  OBLIGACION  -->  PAGO", bold=True, size=14, color=(0, 51, 102), align='center')
doc.add_paragraph()
p("Cada paso depende del anterior. No se puede crear un RP sin un CDP, ni una Obligacion sin un RP, ni un Pago sin una Obligacion.")

# --- 5.1 CDP ---
doc.add_heading('5.1 CDP - Certificado de Disponibilidad Presupuestal', level=2)
ruta('Menu Registro de Movimiento > 1. Disponibilidades (CDP)  /  Boton "1. CDP" en dashboard')

paso(1, "Seleccionar rubro de gasto")
bullet('Se abre la ventana "Buscar Rubro Presupuestal"')
bullet("Escriba en el campo de busqueda (filtra por codigo o nombre)")
bullet("La tabla muestra: Codigo, Cuenta, Apropiacion Definitiva")
bullet("Al seleccionar un rubro, abajo aparece el saldo disponible")
bullet("DOBLE-CLIC en el rubro para seleccionarlo")
nota("Solo se muestran rubros hoja. Si el saldo es $0 o negativo, el sistema muestra error.")

paso(2, "Seleccionar concepto/objeto del gasto")
bullet("Se abre ventana con los conceptos usados anteriormente para ese rubro")
bullet("Opcion A: Haga clic en un concepto de la lista para reutilizarlo")
bullet("Opcion B: Escriba un nuevo concepto en el campo de texto")
bullet('Haga clic en "Seleccionar" o "Usar Nuevo Concepto"')

paso(3, "Seleccionar fuente SIFSE")
bullet('Se abre ventana "Fuente de Financiacion (SIFSE)"')
bullet("Seleccione de la lista desplegable la fuente que financia este gasto")
p("Fuentes mas comunes:", bold=True)
add_table(
    ["Codigo", "Fuente", "Cuando usarla"],
    [
        ["2", "SGP Calidad por Gratuidad (MEN)", "Gastos financiados con recursos de gratuidad"],
        ["1", "Recursos Propios", "Gastos con ingresos de venta de bienes/servicios"],
        ["6", "SGP Matricula - Calidad", "Gastos con recursos de calidad SGP"],
        ["3", "Otras Transferencias", "Gastos con aportes de nacion, departamento o municipio"],
    ],
    col_widths=[1.5, 6, 8]
)
importante("La fuente SIFSE se selecciona UNA sola vez aqui en el CDP. Se propaga automaticamente a RP, Obligacion y Pago.")

paso(4, "Ingresar valor")
bullet("El sistema muestra: rubro, objeto, saldo disponible")
bullet("Ingrese el valor del CDP (solo numeros)")
bullet("El valor no puede exceder el saldo disponible")

paso(5, "Confirmar")
bullet("Revise el resumen: rubro, objeto, fuente SIFSE, valor")
bullet("Haga clic en SI para confirmar")
p("Resultado: Se asigna numero de CDP automaticamente, fecha del dia, estado ACTIVO.", bold=True)

ejemplo("CDP para pago de energia", [
    "Rubro: 2.1.2.02.02.006.04 - Energia electrica",
    "Concepto: Pago factura EMSA febrero 2026",
    "Fuente SIFSE: 2 - SGP Calidad por Gratuidad (MEN)",
    "Valor: $ 800,000",
    "Resultado: CDP No. 1 creado exitosamente",
])

doc.add_page_break()

# --- 5.2 RP ---
doc.add_heading('5.2 RP - Registro Presupuestal (Compromiso)', level=2)
ruta('Menu Registro de Movimiento > 2. Registros Presupuestales (RP)  /  Boton "2. RP"')

paso(1, "Seleccionar CDP")
bullet('Se abre la ventana "Seleccionar CDP Disponible"')
bullet("La tabla muestra: No, Fecha, Rubro, Objeto, Fuente SIFSE, Valor, Saldo, Estado")
bullet("CDPs agotados (saldo $0) aparecen en gris")
bullet("DOBLE-CLIC para seleccionar el CDP")
nota('Desde esta ventana tambien puede Anular CDP (boton rojo) o Editar Valor (boton naranja).')

paso(2, "Ingresar tercero (beneficiario)")
bullet("Ingrese el NIT o cedula del tercero")
bullet("Si no existe, el sistema pregunta si desea crearlo (solo pide nombre)")

paso(3, "Ingresar valor")
bullet("Muestra info del CDP, tercero y saldo disponible")
bullet("Ingrese el valor del RP (no puede exceder el saldo del CDP)")

paso(4, "Ingresar objeto/contrato")
bullet("Escriba la descripcion del compromiso o contrato")

paso(5, "Confirmar")
p("Resultado: RP con numero asignado. La fuente SIFSE se hereda automaticamente del CDP.", bold=True)

ejemplo("RP para pago de energia", [
    "CDP: No. 1 (Energia electrica)",
    "Tercero: 800.456.789-1 - EMSA S.A.",
    "Valor: $ 800,000",
    "Objeto: Contrato suministro energia 2026",
    "Resultado: RP No. 1 creado (fuente SIFSE = 2, heredada del CDP)",
])

# --- 5.3 Obligacion ---
doc.add_heading('5.3 Obligacion', level=2)
ruta('Menu Registro de Movimiento > 3. Obligaciones  /  Boton "3. OBLIGACION"')

paso(1, "Seleccionar RP")
bullet('Ventana "Seleccionar RP" - DOBLE-CLIC para seleccionar')
bullet("Muestra: No, Fecha, CDP, Rubro, Tercero, Valor, Saldo, Estado")

paso(2, "Ingresar valor")
bullet("Puede ser parcial (menor al saldo del RP)")

paso(3, "Factura/Concepto")
bullet("Ingrese numero de factura o concepto de la obligacion")

paso(4, "Confirmar")
p("Resultado: Obligacion registrada. Fuente SIFSE y tercero heredados del RP.", bold=True)

ejemplo("Obligacion por factura de energia", [
    "RP: No. 1 (EMSA S.A.)",
    "Valor: $ 800,000",
    "Factura: FE-2026-00123",
    "Resultado: Obligacion No. 1 creada",
])

doc.add_page_break()

# --- 5.4 Pago ---
doc.add_heading('5.4 Pago', level=2)
ruta('Menu Registro de Movimiento > 4. Pagos  /  Boton "4. PAGO"')

paso(1, "Seleccionar obligacion")
bullet("Ventana con obligaciones pendientes - DOBLE-CLIC para seleccionar")

paso(2, "Ingresar valor")
bullet("Puede ser pago parcial o total")
bullet("Se valida contra el saldo de la obligacion")
bullet("Se valida contra el PAC del mes (si esta configurado)")

paso(3, "Concepto del pago")
bullet("Descripcion breve del pago")

paso(4, "Medio de pago")
bullet("Ventana con dos opciones: SI = Transferencia bancaria, NO = Cheque")

paso(5, "Numero de comprobante de egreso")
bullet("Ingrese el numero del comprobante")

paso(6, "Confirmar")
p("Resultado: Pago registrado con estado PAGADO. Fuente SIFSE heredada.", bold=True)

ejemplo("Pago de factura de energia", [
    "Obligacion: No. 1 (EMSA S.A.)",
    "Valor: $ 800,000",
    "Concepto: Pago factura energia febrero",
    "Medio: Transferencia",
    "Comprobante: CE-001",
    "Resultado: Pago No. 1 registrado",
])

# --- 5.5 Recaudo ---
doc.add_heading('5.5 Recaudo (Ingresos)', level=2)
ruta('Menu Registro de Movimiento > 5. Ingresos / Recaudos  /  Boton "5. RECAUDO"')

paso(1, "Seleccionar rubro de ingreso")
bullet("Ventana de busqueda - busque por codigo o nombre - DOBLE-CLIC")

paso(2, "Concepto")
bullet("Describa el origen del ingreso")

paso(3, "Numero de comprobante de ingreso")

paso(4, "Valor")
bullet("Si excede el presupuesto: muestra advertencia pero permite continuar")

paso(5, "Confirmar")

ejemplo("Recaudo de gratuidad", [
    "Rubro: 1.1.02.06.001.01.03.02 - SGP Calidad por Gratuidad",
    "Concepto: Giro MEN febrero 2026",
    "Comprobante: CI-001",
    "Valor: $ 20,333,039",
    "Resultado: Recaudo No. 1 registrado",
])

doc.add_page_break()

# =====================================================
# 6. MODIFICACIONES
# =====================================================
doc.add_heading('6. Modificaciones Presupuestales', level=1)

doc.add_heading('6.1 Adicion Presupuestal', level=2)
ruta("Menu Registro de Movimiento > 6. Adicion Presupuestal")
p("Aumenta tanto un rubro de gasto como uno de ingreso por el mismo valor. Mantiene el equilibrio.")
numbered("Seleccione el rubro de GASTO a adicionar")
numbered("Seleccione el rubro de INGRESO correspondiente")
numbered("Ingrese el valor de la adicion")
numbered("Ingrese el numero del acto administrativo (Resolucion/Acuerdo)")
numbered("Ingrese una descripcion (opcional)")
numbered("Confirme (se muestra el antes y despues)")

doc.add_heading('6.2 Reduccion Presupuestal', level=2)
ruta("Menu Registro de Movimiento > 7. Reduccion Presupuestal")
p("Disminuye un rubro de gasto y uno de ingreso por el mismo valor. Pasos iguales a la adicion.")
importante("No se puede reducir mas del saldo disponible del rubro de gasto.")

doc.add_heading('6.3 Credito / Contracredito', level=2)
ruta("Menu Registro de Movimiento > 8. Credito / Contracredito")
p("Traslada recursos entre dos rubros de GASTO (no afecta ingresos). El total no cambia.")
numbered("Seleccione el rubro que RECIBE recursos (CREDITO)")
numbered("Seleccione el rubro que CEDE recursos (CONTRACREDITO) - debe ser diferente")
numbered("Ingrese el valor a trasladar (no puede exceder saldo del contracredito)")
numbered("Ingrese numero de acto administrativo y descripcion")
numbered("Confirme")

doc.add_heading('6.4 Ver Modificaciones', level=2)
ruta("Menu Consultas > Ver Modificaciones Presupuestales")
p("Tabla con todas las modificaciones: No, Fecha, Tipo, Acto, Valor, Descripcion.")
p("DOBLE-CLIC en una fila para ver el detalle completo con rubros afectados.", bold=True)

doc.add_page_break()

# =====================================================
# 7. CONSULTAS
# =====================================================
doc.add_heading('7. Consultas y Listados', level=1)

doc.add_heading('7.1 Listados de documentos', level=2)
ruta("Menu Consultas")

add_table(
    ["Opcion de menu", "Que muestra"],
    [
        ["Ver CDPs", "Todos los CDPs con rubro, objeto, fuente SIFSE, valor, estado"],
        ["Ver RPs", "Todos los RPs con CDP origen, tercero, valor, estado"],
        ["Ver Obligaciones", "Todas las obligaciones con RP, tercero, factura, valor"],
        ["Ver Pagos", "Todos los pagos con medio de pago, comprobante, valor"],
        ["Ver Recaudos", "Todos los recaudos por rubro de ingreso"],
    ],
    col_widths=[4.5, 12]
)

p("Acciones disponibles en cada listado:", bold=True)
bullet("Anular (boton rojo): Cambia estado a ANULADO (el registro no se elimina, queda para auditoria)")
bullet("Editar Valor (boton naranja): Permite ajustar el monto del documento")
bullet("Imprimir (boton azul): Genera comprobante HTML para imprimir")
bullet("DOBLE-CLIC en cualquier fila: Abre el comprobante para imprimir")

doc.add_heading('7.2 Tarjeta Presupuestal', level=2)
ruta("Menu Consultas > Tarjeta Presupuestal")
numbered("Seleccione un rubro de gasto")
numbered("Vera todos los movimientos del rubro en orden cronologico")
p("Columnas: Fecha, Tipo, No, NIT, Tercero, Concepto, V.CDP, V.RP, V.Oblig, V.Pago")
bullet("Totales al final de cada columna")
bullet("Saldo disponible resaltado en el pie")
bullet("DOBLE-CLIC en una fila imprime el documento correspondiente")

doc.add_page_break()

# =====================================================
# 8. INFORMES
# =====================================================
doc.add_heading('8. Generacion de Informes', level=1)
ruta("Menu Generacion de Informes > Informes y Ejecuciones")

p("Se abre una ventana con filtros (Mes Inicial y Mes Final) y un arbol de informes disponibles. Seleccione el informe y haga DOBLE-CLIC o presione el boton Generar.")

doc.add_heading('8.1 Ejecucion de Gastos (Formato Catalogo)', level=2)
p("Informe oficial de ejecucion presupuestal de gastos.")
p("Columnas:", bold=True)
bullet("Ppto Inicial, Adiciones, Reducciones, Creditos, Contra-Creditos, Ppto Definitivo")
bullet("Compromisos: Anterior, Mes, Acumulado")
bullet("Pagos: Anterior, Mes, Acumulado")
bullet("Saldo Apropiacion, Saldo Compromisos por Pagar")
bullet("Selector de mes para cambiar el periodo de consulta")
bullet('Boton "Exportar Excel": genera archivo .xlsx con formato profesional, formulas y colores')

doc.add_heading('8.2 Auxiliar - Detalle Movimientos', level=2)
p("Muestra TODOS los movimientos individuales en un periodo (CDPs, RPs, Obligaciones, Pagos), codificados por color segun tipo, con totales por tipo de documento.")

doc.add_heading('8.3 Cadena Presupuestal (CDP > RP > Oblig > Pago)', level=2)
p("Vista jerarquica expandible/colapsable de toda la cadena de ejecucion. Permite ver el flujo completo desde el CDP hasta el pago final.")

doc.add_heading('8.4 Ejecucion de Ingresos (Formato Catalogo)', level=2)
p("Informe oficial con: Ppto Inicial, Adiciones, Reducciones, Ppto Definitivo, Recaudo (Anterior, Mes, Acumulado), Saldo por Recaudar. Exportable a Excel.")

doc.add_heading('8.5 Equilibrio Presupuestal', level=2)
p("Comparativo Ingresos vs Gastos con presupuesto, ejecucion, porcentajes y resultado (Equilibrado / Superavit / Deficit).")

doc.add_page_break()

# =====================================================
# 9. INFORME SIFSE
# =====================================================
doc.add_heading('9. Informe SIFSE (Trimestral)', level=1)
ruta("Menu Generacion de Informes > Informe SIFSE (Trimestral)")
p("Este es el informe que se sube al Sistema de Informacion de Fondos de Servicios Educativos del MEN.")

doc.add_heading('Paso a paso', level=2)

paso(1, "Seleccionar periodo")
add_table(
    ["Trimestre", "Meses que abarca"],
    [
        ["1", "Enero - Febrero - Marzo"],
        ["2", "Abril - Mayo - Junio"],
        ["3", "Julio - Agosto - Septiembre"],
        ["4", "Octubre - Noviembre - Diciembre"],
    ],
    col_widths=[3, 8]
)

paso(2, "Vista previa")
p('La pantalla tiene dos pestanas:')

p('Pestana "Ingresos Presupuestales":', bold=True)
add_table(
    ["Columna", "Descripcion"],
    [
        ["Fuente", "Codigo fuente SIFSE (1, 2, 3, 6, 32...)"],
        ["Descripcion", "Nombre de la fuente"],
        ["Ppto Inicial", "Presupuesto inicial agrupado por fuente"],
        ["Ppto Definitivo", "Presupuesto definitivo agrupado por fuente"],
        ["Recaudado", "Total recaudado en el trimestre"],
    ],
    col_widths=[3.5, 12]
)

p('Pestana "Gastos Presupuestales":', bold=True)
add_table(
    ["Columna", "Descripcion"],
    [
        ["Fuente", "Codigo fuente SIFSE del CDP que origino el gasto"],
        ["Item", "Codigo item de gasto SIFSE (7, 9, 10, 12...)"],
        ["Desc Item", "Descripcion del item SIFSE"],
        ["Ppto Inicial", "Apropiacion inicial agrupada por cruce fuente x item"],
        ["Ppto Definitivo", "Apropiacion definitiva"],
        ["Compromisos", "Total RPs del trimestre"],
        ["Obligaciones", "Total obligaciones del trimestre"],
        ["Pagos", "Total pagos del trimestre"],
    ],
    col_widths=[3.5, 12]
)

paso(3, "Verificar advertencias")
bullet("Si hay rubros SIN mapeo SIFSE, aparece advertencia en amarillo")
bullet('Si hay CDPs sin fuente asignada, aparecen como "N/A"')

paso(4, "Exportar")
numbered('Haga clic en "Exportar .xls"')
numbered("Seleccione ubicacion y nombre del archivo")
numbered("Se genera archivo .xls con 2 hojas:")
bullet("Ingresos_Presupuestales: CODIGO_DANE, ANIO, TRIMESTRE, FUENTE, PPTO_INICIAL, PPTO_DEFINITIVO, RECAUDADO")
bullet("Gastos_Presupuestales: CODIGO_DANE, ANIO, TRIMESTRE, FUENTE, ITEM, PPTO_INICIAL, PPTO_DEFINITIVO, COMPROMISOS, OBLIGACIONES, PAGOS")

importante("Este archivo .xls es el que se sube directamente al portal SIFSE del MEN.")

doc.add_page_break()

# =====================================================
# 10. MAPEO SIFSE
# =====================================================
doc.add_heading('10. Mapeo SIFSE', level=1)
ruta("Menu Configuracion > Mapeo SIFSE")

p("El sistema SIFSE del MEN usa codigos diferentes al catalogo CCPET. Esta pantalla configura la correspondencia entre los codigos del presupuesto y los codigos SIFSE.")

doc.add_heading('Pestana "Ingresos -> Fuente SIFSE"', level=2)
bullet("Muestra todos los rubros de ingreso hoja con: Codigo CCPET, Cuenta, Fuente SIFSE, Descripcion")
bullet("Rubros sin mapear aparecen resaltados en rojo claro")
p("Para editar un mapeo:", bold=True)
numbered("DOBLE-CLIC en el rubro que desea editar")
numbered("Se abre ventana con lista desplegable de fuentes SIFSE")
numbered("Seleccione la fuente correcta")
numbered('Haga clic en "Guardar"')

doc.add_heading('Pestana "Gastos -> Item SIFSE"', level=2)
p("Funciona igual que ingresos. DOBLE-CLIC para editar el item SIFSE de cada rubro de gasto.")

doc.add_heading('Fuentes SIFSE principales', level=2)
add_table(
    ["Cod", "Fuente SIFSE", "Rubros CCPET tipicos"],
    [
        ["1", "Recursos Propios", "1.1.02.05.* (Venta bienes/servicios)"],
        ["2", "SGP Gratuidad", "1.1.02.06.001.01.03.02"],
        ["3", "Otras Transferencias", "Aportes nacion, departamentos, municipios"],
        ["6", "SGP Matricula", "1.1.02.06.006.06.*"],
        ["32", "Rec. Balance Gratuidad", "1.2.10.02.01"],
        ["33", "Rec. Balance Propios", "1.2.10.02.02"],
        ["35", "Rend. Financieros Gratuidad", "1.2.05.02.01"],
        ["36", "Rend. Financieros Otros", "1.2.05.02.02, 1.2.05.02.04"],
        ["52", "Donaciones", "1.2.08.*"],
    ],
    col_widths=[1.2, 5.5, 8.5]
)

doc.add_heading('Items SIFSE principales', level=2)
add_table(
    ["Cod", "Item SIFSE", "Rubros CCPET tipicos"],
    [
        ["7", "Funcionamiento basico", "Aseo, botiquin, otros bienes"],
        ["9", "Servicios publicos", "Acueducto, aseo"],
        ["10", "Energia", "2.1.2.02.02.006.04"],
        ["11", "Telefono", "2.1.2.02.02.008.02"],
        ["12", "Internet / Hosting", "2.1.2.02.02.008.03, .01"],
        ["14", "Seguros", "2.1.2.02.02.007.01"],
        ["15", "Serv. profesionales", "2.1.2.02.02.008.05, .06"],
        ["20", "Mant. infraestructura", "2.1.2.02.02.005.01, .008.07"],
        ["21", "Dotacion institucional", "Maquinaria, muebles, software"],
        ["22", "Material pedagogico", "Papeleria, deporte, dotacion"],
        ["23", "Transporte escolar", "2.1.2.02.02.006.06"],
        ["25", "Alimentacion escolar", "2.1.2.02.02.006.07"],
        ["26", "Act. pedagogicas", "2.1.2.02.02.009.*"],
        ["86", "Servicios financieros", "2.1.2.02.02.007.02"],
    ],
    col_widths=[1.2, 5, 8.5]
)

p('Boton "Restaurar Mapeo por Defecto": borra los mapeos actuales y aplica los predeterminados.', bold=True)

doc.add_page_break()

# =====================================================
# 11. PAC
# =====================================================
doc.add_heading('11. Plan Anualizado de Caja (PAC)', level=1)
ruta('Menu Plan Presupuestal > Plan Anualizado de Caja  /  Boton "PAC" en dashboard')

p("El PAC distribuye la apropiacion de cada rubro de gasto en los 12 meses del ano. Si esta configurado, el sistema VALIDA que cada pago no exceda el cupo mensual programado.")

doc.add_heading('Pantalla PAC', level=2)
p("Panel izquierdo: Lista de rubros de gasto", bold=True)
bullet("Busqueda por codigo o nombre")
bullet("Seleccione un rubro para ver/editar su PAC")

p("Panel derecho: Distribucion mensual", bold=True)
bullet("12 campos editables (uno por mes)")
bullet("Cada fila muestra: Mes, Programado, Ejecutado (pagos reales), Disponible")
bullet("Total al final")
bullet("Advertencia si el total no coincide con la apropiacion definitiva")

p("Botones:", bold=True)
bullet("Distribuir Uniforme: Divide la apropiacion en 12 partes iguales automaticamente")
bullet("Guardar PAC: Guarda la distribucion ingresada manualmente")

nota("Si el PAC no esta configurado para un rubro, los pagos de ese rubro no tienen restriccion mensual.")

# =====================================================
# 12. PROCESOS
# =====================================================
doc.add_heading('12. Procesos Presupuestales', level=1)

doc.add_heading('12.1 Consolidar Mes', level=2)
ruta("Menu Procesos Presupuestales > Consolidar Mes")
p("Agrega los movimientos del mes actual para cada rubro: compromisos (RPs), pagos y recaudos (ingresos).")
nota("Ejecute esto al final de cada mes antes del cierre.")

doc.add_heading('12.2 Cierre de Mes', level=2)
ruta("Menu Procesos Presupuestales > Cierre de Mes")
numbered("Confirme el cierre")
numbered("El sistema consolida automaticamente")
numbered("Avanza el mes activo al siguiente")
numbered('Si cierra diciembre, muestra "FIN DE VIGENCIA"')
importante("El cierre de mes no es reversible. Asegurese de que todos los movimientos del mes esten registrados antes de cerrar.")

doc.add_page_break()

# =====================================================
# 13. BACKUPS
# =====================================================
doc.add_heading('13. Copias de Seguridad', level=1)

doc.add_heading('13.1 Crear Backup', level=2)
ruta("Menu Configuracion > Crear Copia de Seguridad")
numbered("Seleccione la carpeta destino")
numbered("Se crea archivo: presupuesto_ie_2026_backup_AAAAMMDD_HHMMSS.db")
numbered("Muestra el tamano del archivo y confirma")
importante("Haga backup frecuentemente, especialmente antes de modificaciones presupuestales importantes.")

doc.add_heading('13.2 Restaurar Backup', level=2)
ruta("Menu Configuracion > Restaurar Copia de Seguridad")
numbered("ADVERTENCIA: Esto reemplaza TODOS los datos actuales")
numbered("Seleccione el archivo .db de respaldo")
numbered("Confirme la restauracion")
numbered("El sistema se reinicia automaticamente")

doc.add_page_break()

# =====================================================
# 14. GUIA DE PRUEBAS
# =====================================================
doc.add_heading('14. Guia Rapida de Pruebas', level=1)
p("Siga estos pasos en orden para probar todo el sistema. Tiempo estimado total: 40 minutos.", bold=True, size=12)

# Prueba 1
doc.add_heading('PRUEBA 1: Configuracion (5 min)', level=2)
numbered('Abra Configuracion > Configuracion General')
numbered("Llene todos los campos: institucion, NIT, DANE, rector, tesorera")
numbered("Verifique que el Mes Actual sea 1 (Enero) o el mes correcto")
numbered("Guarde")
numbered("Verifique que el dashboard muestre los datos correctos en el encabezado")

doc.add_heading('PRUEBA 2: Catalogo presupuestal (ya importado)', level=2)
numbered("Revise Plan Presupuestal > Rubros de Gastos")
numbered('Busque un rubro conocido (ej: "energia")')
numbered("Verifique que aparezca con la apropiacion correcta")
numbered("Revise Rubros de Ingresos de la misma forma")
numbered("Verifique el Resumen Presupuestal - debe mostrar equilibrio")

doc.add_heading('PRUEBA 3: Registrar un tercero (2 min)', level=2)
numbered("Vaya a Gestion de Terceros > Registrar Tercero")
numbered("Ingrese: NIT=900123456, DV=7, Nombre=EMPRESA DE ENERGIA SA, Tipo=Juridica")
numbered("Guarde y verifique en Ver Terceros")

doc.add_heading('PRUEBA 4: Registrar CDP completo (3 min)', level=2)
numbered('Haga clic en "1. CDP" en el dashboard')
numbered('Busque el rubro "Energia" (ej: 2.1.2.02.02.006.04)')
numbered("Doble-clic para seleccionar")
numbered('Escriba concepto: "Pago energia electrica enero 2026"')
numbered("Seleccione fuente SIFSE: 2 - SGP Calidad por Gratuidad (MEN)")
numbered("Ingrese valor: 500000")
numbered("Confirme")
p("Resultado esperado: CDP No. 1 creado exitosamente, fuente SIFSE = 2", bold=True, color=(0, 102, 51))

doc.add_heading('PRUEBA 5: Registrar RP (3 min)', level=2)
numbered('Haga clic en "2. RP"')
numbered("En la ventana de CDPs, verifique que aparece el CDP recien creado con columna Fuente = 2")
numbered("Doble-clic en el CDP No. 1")
numbered("Ingrese NIT: 900123456 (el tercero creado antes)")
numbered("Valor: 500000")
numbered('Objeto: "Contrato energia enero 2026"')
numbered("Confirme")
p("Resultado esperado: RP No. 1, fuente SIFSE heredada automaticamente del CDP", bold=True, color=(0, 102, 51))

doc.add_heading('PRUEBA 6: Registrar Obligacion (2 min)', level=2)
numbered('Haga clic en "3. OBLIGACION"')
numbered("Doble-clic en el RP No. 1")
numbered("Valor: 500000")
numbered('Factura: "FE-2026-001"')
numbered("Confirme")
p("Resultado esperado: Obligacion No. 1 creada", bold=True, color=(0, 102, 51))

doc.add_heading('PRUEBA 7: Registrar Pago (2 min)', level=2)
numbered('Haga clic en "4. PAGO"')
numbered("Doble-clic en la Obligacion No. 1")
numbered("Valor: 500000")
numbered('Concepto: "Pago factura energia enero"')
numbered("Medio: SI (Transferencia)")
numbered('Comprobante: "CE-001"')
numbered("Confirme")
p("Resultado esperado: Pago No. 1 registrado", bold=True, color=(0, 102, 51))

doc.add_heading('PRUEBA 8: Registrar Recaudo (2 min)', level=2)
numbered('Haga clic en "5. RECAUDO"')
numbered("Busque el rubro de ingreso de Gratuidad (1.1.02.06.001.01.03.02)")
numbered('Concepto: "Giro MEN primer trimestre 2026"')
numbered('Comprobante: "CI-001"')
numbered("Valor: 20333039")
numbered("Confirme")
p("Resultado esperado: Recaudo No. 1 registrado", bold=True, color=(0, 102, 51))

doc.add_heading('PRUEBA 9: Verificar dashboard (1 min)', level=2)
p("Regrese al dashboard y verifique:")
bullet("Total CDPs: $ 500,000")
bullet("Comprometido (RPs): $ 500,000")
bullet("Total Obligado: $ 500,000")
bullet("Total Pagado: $ 500,000")
bullet("En ingresos: Total Recaudado: $ 20,333,039")

doc.add_heading('PRUEBA 10: Consultas (3 min)', level=2)
numbered("Consultas > Ver CDPs: debe mostrar CDP No. 1 con fuente SIFSE")
numbered("Consultas > Ver RPs: debe mostrar RP No. 1")
numbered("Consultas > Tarjeta Presupuestal: seleccione rubro de energia")
p("Debe mostrar toda la cadena: CDP, RP, Obligacion, Pago en la tarjeta.", italic=True)

doc.add_heading('PRUEBA 11: Informes (5 min)', level=2)
numbered("Informes > Informes y Ejecuciones")
numbered("Seleccione Ejecucion de Gastos > doble-clic")
numbered("Verifique que el rubro de energia muestra los montos correctos")
numbered("Pruebe el boton Exportar Excel")
numbered("Repita con Ejecucion de Ingresos")
numbered("Revise Cadena Presupuestal: debe mostrar el arbol CDP>RP>Oblig>Pago")
numbered("Revise Equilibrio Presupuestal")

doc.add_heading('PRUEBA 12: Informe SIFSE (5 min) - LA MAS IMPORTANTE', level=2)
numbered("Informes > Informe SIFSE (Trimestral)")
numbered("Trimestre: 1, Ano: 2026")
numbered("Haga clic en Vista Previa")
numbered("Pestana Ingresos: Verifique Fuente 2 (Gratuidad) con el recaudo de $20,333,039")
numbered("Pestana Gastos: Verifique Fuente 2 + Item 10 (Energia) con los montos de $500,000")
numbered("Haga clic en Exportar .xls")
numbered("Abra el archivo generado en Excel y verifique las 2 hojas")
p("Este archivo es el que se sube al portal SIFSE del MEN.", bold=True, color=(192, 0, 0))

doc.add_heading('PRUEBA 13: Mapeo SIFSE (3 min)', level=2)
numbered("Configuracion > Mapeo SIFSE")
numbered("Pestana Ingresos: verifique que los rubros tienen fuente asignada")
numbered("Rubros sin mapeo aparecen en rojo claro")
numbered("Doble-clic en un rubro para cambiar su fuente - seleccione otra - Guardar")
numbered("Pestana Gastos: verificar items asignados")
numbered('Probar boton "Restaurar Mapeo por Defecto"')

doc.add_heading('PRUEBA 14: Modificacion presupuestal (3 min)', level=2)
numbered('Haga clic en "6. ADICION" en el dashboard')
numbered("Seleccione un rubro de gasto y uno de ingreso")
numbered('Valor: 100000, Acto: "Resolucion 001"')
numbered("Confirme y verifique que las apropiaciones cambiaron")
numbered("Revise en Consultas > Ver Modificaciones")

doc.add_heading('PRUEBA 15: Backup (2 min)', level=2)
numbered("Configuracion > Crear Copia de Seguridad")
numbered("Seleccione una carpeta")
numbered("Verifique que el archivo .db se creo correctamente")

doc.add_page_break()

# =====================================================
# ATAJOS Y TIPS
# =====================================================
doc.add_heading('Resumen de Atajos y Tips', level=1)
add_table(
    ["Accion", "Como hacerlo"],
    [
        ["Seleccionar en cualquier lista", "DOBLE-CLIC en la fila"],
        ["Buscar rubros", "Escriba codigo o nombre en el campo de busqueda"],
        ["Imprimir comprobante", "Doble-clic en listado o boton Imprimir"],
        ["Editar mapeo SIFSE", "Doble-clic en rubro (pantalla Mapeo SIFSE)"],
        ["Ver detalle de modificacion", "Doble-clic en listado de modificaciones"],
        ["Exportar informe a Excel", "Boton Exportar Excel en pantalla del informe"],
        ["Volver al inicio", "Boton < Volver o clic en boton del dashboard"],
    ],
    col_widths=[6, 10]
)

doc.add_heading('Flujo Recomendado Mensual', level=1)

p("1. Inicio de mes:", bold=True)
bullet("Verificar mes activo en Configuracion")

p("2. Durante el mes:", bold=True)
bullet("Registrar CDPs segun necesidad (con fuente SIFSE)")
bullet("Crear RPs al comprometer con proveedores")
bullet("Registrar Obligaciones al recibir facturas")
bullet("Registrar Pagos al realizar transferencias o cheques")
bullet("Registrar Recaudos al recibir ingresos")

p("3. Fin de mes:", bold=True)
bullet("Verificar saldos en el dashboard")
bullet("Generar informe de ejecucion de gastos e ingresos")
bullet("Consolidar mes (Procesos > Consolidar)")
bullet("Cierre de mes (Procesos > Cierre)")
bullet("Crear copia de seguridad")

p("4. Fin de trimestre (adicional):", bold=True)
bullet("Generar Informe SIFSE (Informes > Informe SIFSE Trimestral)")
bullet("Exportar archivo .xls")
bullet("Subir al portal SIFSE del MEN")

doc.add_paragraph()
doc.add_paragraph()
p("Manual generado para el Sistema Presupuestal IE 2026", italic=True, size=9, align='center')
p("IE Almirante Padilla - DANE 186755000015", italic=True, size=9, align='center')

# ============ GUARDAR ============
destino = os.path.join(os.path.dirname(__file__), "MANUAL_OPERACION_SISTEMA_PRESUPUESTAL.docx")
doc.save(destino)
print(f"Manual generado exitosamente: {destino}")
